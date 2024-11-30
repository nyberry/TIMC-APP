import os
import docx
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement

from word_generator import modify_table as tab


def write_word_document(document, patient, user, workbook, data):

    print (f'Start writing the document of type: {document.type}')
    sheet = None            #sheet will be set later only if we plan to include spreadsheet data

    #function to search for the markers of the start and end of the bloods fields:
    def find_start_row(sheet):
        for cell in sheet.range('B1',"B1000"):
            if cell.value == 'Blood tests':
                return(cell.row)

    #function to call to delete unwanted paragraphs
    def delete_paragraph(doc, index):
        if index < 0 or index >= len(doc.paragraphs):
            print(f"Invalid paragraph index: {index}")
            return
        paragraph = doc.paragraphs[index]
        p_element = paragraph._element
        body_element = p_element.getparent()
        body_element.remove(p_element)
        return

    # Depending on the document type, define the header layout and the parts to keep
    # paragraphs are: 0: title; 1: horizontal line; 2: name; 3: qid; 4: clinic ref; 5: phone, 6: date; 7: doctor; 8: Recipient; 9: Sender
    
    if document.type == 'Medical report':
        if patient.medas_attachments:
            document.description = f"Medical report for {document.patient}"
            document.title = "Medical report"
            paragraphs_to_keep = [0,1,2,4,6,7]
        else:
            document.title = "Test report"
            document.description = f"Test report for {document.patient}"
            paragraphs_to_keep = [0,1,2,4,6]
    
    elif document.type == "Letter with lab results":
        document.title = "Test results"
        document.description = f"Lab results for {document.patient}"
        paragraphs_to_keep = [0,1,2,4,6]

    elif document.type == 'Letter of referral':
        document.title = "Referral letter"
        document.description = f"Referral letter for {document.patient}"
        paragraphs_to_keep = [0,1,2,3,4,5,6,8,9]  

    elif document.type == 'Letter to patient':
        document.title = "Medical letter"
        document.description = f"Letter to {document.patient}"
        paragraphs_to_keep = [0,1,2,4,6]
        
    elif document.type == 'Letter TWIMC':
        document.title = "Medical letter"
        document.description = f"Medical letter for {document.patient}"
        paragraphs_to_keep = [0,1,2,4,6]

    else:
        document.type == 'Unspecified document'
        document.title = "Medical document"
        document.description = f"Medical document for {document.patient}"
        paragraphs_to_keep = [0,1,2,3,4,5,6,7,8,9]

        
    # Open the Word document and Excel Sheet
    if document.type == 'send by whatsapp':
        doc = docx.Document()
        document.title = "WhatsApp attachment"
        document.description = f"WhatsApp attachment {document.patient}"
        # doc.add_paragraph(document.description)
        print ("empty doc object created but not displayed")
        return doc
    
    else:
        template_filepath = os.path.abspath("../data/HStemplate.docx")
        doc = docx.Document(template_filepath)

    # Add the details to the header
    try:
        # Add title
        doc.paragraphs[0].clear()
        doc.paragraphs[0].add_run(document.title)

        # Add name
        doc.paragraphs[2].add_run(f'{patient.title} {patient.firstname} {patient.surname}')

        # QID
        doc.paragraphs[3].add_run(patient.QID)

        # Clinic ref
        if 'Clinic reference' in data:                                                                
            doc.paragraphs[4].add_run(data["Clinic reference"])
      
        # phone
        doc.paragraphs[5].add_run(f" {patient.phone}")

        # add date
        today_date = datetime.now().date()
        today_date = today_date.strftime("%dth %B %Y")
        doc.paragraphs[6].add_run(f': {today_date}')

        # doctor
        if 'Doctor' in data:
            doc.paragraphs[7].add_run(f" {data['Doctor']}")
         
        # to and from
        doc.paragraphs[9].add_run(document.author)
       
        # delete unwanted paragraphs
     
        for paragraph in range(9,-1,-1):
            if paragraph not in paragraphs_to_keep:
                delete_paragraph(doc, paragraph)

    except Exception as e:
        print (f"Error writing header {e}")


    # Write body of document
    if  document.title=="Test results":
        doc.add_paragraph(f'Dear {patient.firstname},')
        doc.add_paragraph("Attached are the results of the recent tests. If there is anything you would like to discuss, please let me know.")

    elif document.type=="Letter of referral":
        try:
            doc.add_paragraph(f'Reason for referral:').bold=True
            doc.add_paragraph(f"Thankyou for seeing {patient.firstname} who is aged {data['Age']} and is usually in good health. {data['Pronoun'].capitalize()} consulted because...")
            doc.add_paragraph(f'On examination ...')
            doc.add_paragraph(f'Blood tests show...')
            doc.add_paragraph(f'My concern is that...')
            doc.add_paragraph(f'Thankyou for your help with...')
        except Exception as e:
            print (f"problem adding content to referral letter {e}")
    
    elif document.type=="Letter to patient":
        doc.add_paragraph(f'Dear {patient.firstname},')
        doc.add_paragraph(f'Text of letter')

    elif document.type=="Letter TWIMC":
        doc.add_paragraph(f'Dear Sir or Madam,')
        doc.add_paragraph(f'Text of letter')

    elif document.title=="Test report":
        # Add a brief intro 
        doc.add_paragraph(f'Dear {patient.firstname},')
        doc.add_paragraph("This is a summary of the results of the tests we arranged recently.")
        # find the spreadsheet row to start adding blood test explanations from:
        sheet = workbook.sheets['Text']
        text_start_row = find_start_row(sheet)

    elif document.title=="Medical report":
        # Add a warm intro
        doc.add_paragraph(f'Dear {patient.firstname}, ')
        doc.add_paragraph("It was nice meeting you recently for a health check. This is a summary of our discussion and findings, with some recommendations.")
        text = "By way of background, you are aged " + str(int(data['Age']))+ ". "
        if 'Occupation' in data:
            text+=str("You work in "+data['Occupation']+". ")
        if 'Reason for health check' in data:
            text+=str("You arranged this health check "+data['Reason for health check']+". ")  
        doc.add_paragraph(text)

        # Past Medical history
        if 'Past Medical History' in data:
            text = "Thankyou for telling me about your health over the years. "
            if 'nil' in data['Past Medical History'].lower():
                text+="You are generally in good health, and have no significant medical problems. You have had no major operations. Your have had no significant illnesses or injuries requiring hospitalisation. "
            else:
                text+="You are generally in good health, and have had few significant medical problems. Previously, you have been identified to have "+data['Past Medical History']+". "
                text+="Otherwise, you have had no significant illnesses or injuries requiring hospitalisation."
            doc.add_paragraph(text)
        if 'known diabetes' in data:    
            if data['Known diabetes?']=="":
                data['Known diabetes?']="no"
        if 'known high cholesterol' in data:
            if data['Known high cholesterol?']=="":
                data['Known high cholesterol?']="no"
        if 'known hypertension' in data:
            if data['Known hypertension?']=="":
                data['Known hypertension?']="no"
        
        #allergies 
        allergytext=""
        if 'Allergies' in data:
            if data['Allergies'] == "nil":
                allergytext="You have no significant allergies. "
            elif data['Allergies'] == "NKDA":
                allergytext="You are not allergic to any medications. "
            elif data['Allergies'] == 'unknown':
                allergytext="You have no known allergies. "
            else:
                allergytext="In the past you have had allergy or sensitivity to "+data['Allergies']+". "
            doc.add_paragraph(allergytext)
            
        #meds
        if 'Medications' in data:
            if data['Medications']== "nil" or data['Medications']=="":
                text="You do not regularly take any medications. "+allergytext
                doc.add_paragraph(text)
            else:
                meds=data['Medications'].split(',')
                meds=[med.strip() for med in meds]
                doc.add_paragraph("Medications:")
                for i in range(len(meds)):
                    paragraph=doc.add_paragraph(str(i+1)+". "+meds[i])
                    paragraph.paragraph_format.left_indent = Cm(1)
                doc.add_paragraph(allergytext)

        #family history
        if 'Family History' in data:
            if data['Family History']=='nil':
                text="In your immediate family, there is no history of significant illness such as diabetes, heart disease, cancer or stroke. There is no history of inherited or genetic conditions, to your knowledge." 
            elif data['Family History']!="" :
                text="In your family, you mentioned that there is a history of "+data['Family History']
            doc.add_paragraph(text)

        #lifestyle
        text = ""
        if 'Smoking' in data:
            text+="You "+data['Smoking']+". "
        if 'Alcohol' in data:
            text+="You "+data['Alcohol']+". "
        #if 'Alcohol units' in data:
        #    text+data['Alcohol units']+". "
        if 'Exercise' in data:
            text+="You "+data['Exercise']
        if text != "":
            doc.add_paragraph(text)

        #cervical smears
        if 'Cervical smears' in data:
            if data['Cervical smears'] != "":
                text = "We discussed your previous cervical smears, and I recorded that "+data['Cervical smears']
                doc.add_paragraph(text)

        #history and systems enquiry
        if 'History' in data:
            doc.add_paragraph(data['History'])
        if 'Systems enquiry' in data:
            if data['Systems enquiry']=='nil':
                text="Systematic enquiry did not reveal any other symptoms of concern. "
            elif data['Systems enquiry']!="" :
                text="You mentioned experiencing "+data['Systems enquiry']
            doc.add_paragraph(text)
        if 'Mood screening' in data:
            if data['Mood screening'] == 'nil':
                text='You described your mood as good and stable. There were no pointers to anxiety or depression.'
            else:
                text=data['Mood screening']
            doc.add_paragraph(text)

        # prime to add data from spreadsheet
        text_start_row = 1     #for the next part of the report, we will start at the top
        sheet = workbook.sheets['Text']

    # thats all we need to do if just sending attachments by whatsapp
    elif document.type == "send by whatsapp":
        print("Empty doc generated")
        return doc

    # there has been an error and the report type is not defined
    else:
        print("Error: report type not specified")
        text_start_row = 1
        if workbook:
            sheet = workbook.sheets['Text']


    # Add data from spreadsheet if appropriate
    if sheet:
        row = text_start_row
        while sheet.range(f'F{row}').value != 'END':
            cell = sheet.range(f'F{row}')
            text = cell.value
            if text!=None:
                text = str(cell.value)
                if text == "(recommendations)":
                    handle_recommendations(sheet,doc,text_start_row)
                elif cell.api.Font.Bold:
                    doc.add_heading(text, 1)   #interpret bold cells on the spreadsheet as titles
                elif cell.api.Font.Italic:
                    paragraph=doc.add_paragraph(u'\u2022'+"  "+text)  # this adds a unicode bullet point'
                    paragraph.paragraph_format.left_indent = Cm(0.7)  # and indents it a little  
                else:
                    doc.add_paragraph(text)
            row+=1


    # Sign off
    doc.add_paragraph("Yours,")
    doc.add_paragraph()
    doc.add_paragraph(document.author)


    # Add stamps (they can always be deleted by user)
    try:
        # Add personal stamp if exists
        table = doc.add_table(rows=1, cols=2)
        if user.stamp:
            cell = table.cell(0, 0)
            cell_paragraph = cell.paragraphs[0]
            run = cell_paragraph.add_run()
            image_filepath = os.path.abspath(f"../images/stamps/{user.stamp}")
            run.add_picture(image_filepath, width = Cm(6))

        # Add the TIMC stamp to the second cell
        cell = table.cell(0, 1)
        cell_paragraph = cell.paragraphs[0]
        run = cell_paragraph.add_run()
        image_filepath = os.path.abspath("../images/stamps/TIMC_stamp.jpg")
        run.add_picture(image_filepath, width = Cm(4))
    
    except Exception as e:
        print (f"Error adding stamps {e}")


    # add disclaimer if needed
    if document.type == "Medical report":
        try:
            disclaimer_message = f'Disclaimer: This document is intended for the purpose of conveying health information. It is confidential and exclusively for the person named. It is based on available medical records, clinical assessments, and diagnostic tests. Additional medical data may be necessary for a comprehensive understanding of a person\'s health. It may not cover all possible medical conditions, and there may be other factors not addressed here. '
            disclaimer_message += f'The information presented was checked for accuracy, but errors or inaccuracies may still sometimes occur. Advice is based on the current state of medical knowledge and may be subject to change. The International Medical Centre takes no responsibility for the use or interpretation of this document by other individuals or entities. '
            disclaimer_message += f'For questions or clarifications, please contact The International Medical Centre Tel: +974 4488 429 / +974 6644 4282, Email: info@theimcentre.com'
            doc.add_paragraph(disclaimer_message)
            print ("Disclaimer added")          
        except Exception as e:
            print ("Unable to add a disclaimer {e}")   


    # Add name to footer
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    custom_text = f"TIMC: {patient.firstname} {patient.surname}"
    footer_paragraph.add_run(custom_text)

    return (doc)

def handle_recommendations(sheet,doc,text_start_row):
# iterate through the cells in the recommendation column and add each cell's content as a bullet point paragraph, until reach END flag   
    row = text_start_row
    while sheet.range(f'G{row}').value != 'END':
        cell = sheet.range(f'G{row}')
        text = cell.value
        if text!=None:
            text = str(cell.value)
            paragraph=doc.add_paragraph(u'\u2022'+"  "+text)  # this adds a unicode bullet point'
            paragraph.paragraph_format.left_indent = Cm(0.5)  # and indents it a little  
                  #interpret bold cells on the spreadsheet as titles
        row+=1
